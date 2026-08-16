"""DOCX parser (§3.2.7, FR-KB-1). Unlike PDF, DOCX carries real semantic
structure — paragraph styles ("Heading 1".."Heading 6" vs "Normal") and
native table elements — so this parser recovers genuine heading levels
and whole-table nodes, not a heuristic. Walks the document body's XML
children directly (rather than ``doc.paragraphs``/``doc.tables``
separately) so headings, paragraphs, and tables come out in their real
reading order, not paragraphs-then-tables.
"""

from __future__ import annotations

import io
import zipfile

from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from aether.app.ingestion.document_tree import DocumentNode, NodeKind

_HEADING_PREFIX = "Heading "


class DocxParseError(Exception):
    """The file is not a valid DOCX package (corrupt, wrong format
    despite passing magic-byte detection, etc.)."""


def parse_docx(content: bytes) -> list[DocumentNode]:
    try:
        document = Document(io.BytesIO(content))
    except (PackageNotFoundError, zipfile.BadZipFile) as exc:
        # BadZipFile: python-docx only wraps *some* malformed-package
        # cases in its own PackageNotFoundError — a file that isn't a
        # zip at all surfaces the underlying zipfile error directly.
        raise DocxParseError(str(exc)) from exc

    nodes: list[DocumentNode] = []
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            node = _paragraph_node(Paragraph(child, document))
            if node is not None:
                nodes.append(node)
        elif child.tag == qn("w:tbl"):
            node = _table_node(Table(child, document))
            if node is not None:
                nodes.append(node)

    if not nodes:
        raise DocxParseError("no extractable content (empty document)")
    return nodes


def _paragraph_node(paragraph: Paragraph) -> DocumentNode | None:
    text = paragraph.text.strip()
    if not text:
        return None
    style_name = paragraph.style.name if paragraph.style is not None else ""
    if style_name.startswith(_HEADING_PREFIX):
        level_str = style_name.removeprefix(_HEADING_PREFIX).strip()
        level = int(level_str) if level_str.isdigit() else 1
        return DocumentNode(
            kind=NodeKind.HEADING, text=text, level=min(max(level, 1), 6), page=None
        )
    return DocumentNode(kind=NodeKind.PARAGRAPH, text=text, level=None, page=None)


def _table_node(table: Table) -> DocumentNode | None:
    rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
    rows = [row for row in rows if any(cell for cell in row)]
    if not rows:
        return None
    text = "\n".join(" | ".join(row) for row in rows)
    return DocumentNode(kind=NodeKind.TABLE, text=text, level=None, page=None)
