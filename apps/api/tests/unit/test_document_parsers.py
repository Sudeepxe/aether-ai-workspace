from __future__ import annotations

import io

import docx
import pytest
from pypdf import PdfWriter

from aether.app.ingestion.document_tree import NodeKind
from aether.app.ingestion.parsers.docx_parser import DocxParseError, parse_docx
from aether.app.ingestion.parsers.html_parser import HtmlParseError, parse_html
from aether.app.ingestion.parsers.markdown_parser import MarkdownParseError, parse_markdown
from aether.app.ingestion.parsers.pdf_parser import PdfParseError, parse_pdf
from aether.app.ingestion.parsers.text_parser import TextParseError, parse_text

pytestmark = pytest.mark.unit


# --------------------------------------------------------------------- PDF --


def _real_pdf(*, pages: list[str]) -> bytes:
    from pypdf import PageObject
    from pypdf.generic import RectangleObject

    writer = PdfWriter()
    for text in pages:
        page = PageObject.create_blank_page(writer, width=612, height=792)
        # pypdf's writer has no simple "add text" helper — build a
        # minimal content stream directly so extract_text() has real
        # text to recover, proving round-trip extraction genuinely
        # works rather than asserting against an empty page.
        content = f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET".encode()
        from pypdf.generic import ContentStream, DictionaryObject, NameObject

        stream = ContentStream(None, writer)
        stream.set_data(content)
        page[NameObject("/Contents")] = writer._add_object(stream)
        page[NameObject("/Resources")] = DictionaryObject(
            {
                NameObject("/Font"): DictionaryObject(
                    {
                        NameObject("/F1"): writer._add_object(
                            DictionaryObject(
                                {
                                    NameObject("/Type"): NameObject("/Font"),
                                    NameObject("/Subtype"): NameObject("/Type1"),
                                    NameObject("/BaseFont"): NameObject("/Helvetica"),
                                }
                            )
                        )
                    }
                )
            }
        )
        page[NameObject("/MediaBox")] = RectangleObject((0, 0, 612, 792))
        writer.add_page(page)
    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


def test_pdf_parser_extracts_text_per_page_with_correct_page_numbers() -> None:
    pdf_bytes = _real_pdf(pages=["First page content", "Second page content"])

    nodes = parse_pdf(pdf_bytes)

    pages_seen = {n.page for n in nodes}
    assert pages_seen == {1, 2}
    assert all(n.kind == NodeKind.PARAGRAPH for n in nodes)
    combined = " ".join(n.text for n in nodes)
    assert "First page" in combined
    assert "Second page" in combined


def test_pdf_parser_rejects_a_corrupt_file() -> None:
    with pytest.raises(PdfParseError):
        parse_pdf(b"not a real pdf at all, just garbage bytes")


def test_pdf_parser_rejects_a_page_with_no_extractable_text() -> None:
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    buf = io.BytesIO()
    writer.write(buf)

    with pytest.raises(PdfParseError):
        parse_pdf(buf.getvalue())


# -------------------------------------------------------------------- DOCX --


def test_docx_parser_recovers_headings_paragraphs_and_tables_in_order() -> None:
    document = docx.Document()
    document.add_heading("Title", level=1)
    document.add_paragraph("Intro paragraph.")
    document.add_heading("Subsection", level=2)
    document.add_paragraph("Body paragraph.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "C"
    table.cell(1, 1).text = "D"
    buf = io.BytesIO()
    document.save(buf)

    nodes = parse_docx(buf.getvalue())

    kinds = [n.kind for n in nodes]
    assert kinds == [
        NodeKind.HEADING,
        NodeKind.PARAGRAPH,
        NodeKind.HEADING,
        NodeKind.PARAGRAPH,
        NodeKind.TABLE,
    ]
    assert nodes[0].text == "Title"
    assert nodes[0].level == 1
    assert nodes[2].text == "Subsection"
    assert nodes[2].level == 2
    assert "A" in nodes[4].text and "D" in nodes[4].text


def test_docx_parser_rejects_a_non_docx_file() -> None:
    with pytest.raises(DocxParseError):
        parse_docx(b"this is not a docx file")


# ----------------------------------------------------------------- Markdown --


def test_markdown_parser_recovers_heading_levels_and_strips_inline_markup() -> None:
    md = b"# Title\n\nIntro with **bold** text.\n\n## Sub\n\nMore body text.\n"

    nodes = parse_markdown(md)

    assert nodes[0].kind == NodeKind.HEADING
    assert nodes[0].level == 1
    assert nodes[0].text == "Title"
    assert nodes[1].kind == NodeKind.PARAGRAPH
    assert nodes[1].text == "Intro with bold text."  # markup stripped, not raw source
    assert nodes[2].kind == NodeKind.HEADING
    assert nodes[2].level == 2


def test_markdown_parser_rejects_empty_content() -> None:
    with pytest.raises(MarkdownParseError):
        parse_markdown(b"   \n\n   ")


# --------------------------------------------------------------------- HTML --


def test_html_parser_recovers_headings_paragraphs_and_tables_in_order() -> None:
    html = (
        b"<html><body>"
        b"<h1>Title</h1><p>Intro.</p>"
        b"<table><tr><td>A</td><td>B</td></tr></table>"
        b"<h2>Sub</h2><p>Body.</p>"
        b"</body></html>"
    )

    nodes = parse_html(html)

    kinds = [n.kind for n in nodes]
    assert kinds == [
        NodeKind.HEADING,
        NodeKind.PARAGRAPH,
        NodeKind.TABLE,
        NodeKind.HEADING,
        NodeKind.PARAGRAPH,
    ]
    assert nodes[0].level == 1
    assert nodes[3].level == 2
    assert "A | B" in nodes[2].text


def test_html_parser_does_not_double_count_paragraphs_inside_tables() -> None:
    html = b"<html><body><table><tr><td><p>cell text</p></td></tr></table></body></html>"

    nodes = parse_html(html)

    assert len(nodes) == 1
    assert nodes[0].kind == NodeKind.TABLE


def test_html_parser_rejects_content_with_no_structural_elements() -> None:
    with pytest.raises(HtmlParseError):
        parse_html(b"<html><body><span>just inline text, no p/h/table</span></body></html>")


# --------------------------------------------------------------------- Text --


def test_text_parser_splits_on_blank_lines() -> None:
    text = b"First paragraph.\n\nSecond paragraph.\n\nThird."

    nodes = parse_text(text)

    assert len(nodes) == 3
    assert all(n.kind == NodeKind.PARAGRAPH for n in nodes)
    assert nodes[0].text == "First paragraph."


def test_text_parser_rejects_empty_content() -> None:
    with pytest.raises(TextParseError):
        parse_text(b"   \n\n  ")
